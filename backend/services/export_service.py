from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from datetime import datetime

class ExportService:
    def export_txt(self, video):
        """Export video to TXT format"""
        lines = [
            "Video Transcription & Analysis",
            "=" * 50,
            "",
            f"Video URL: {video.get('video_url', 'N/A')}",
            f"Platform: {video.get('platform', 'unknown')}",
            f"Duration: {video.get('duration_minutes', 0):.1f} minutes",
            f"Processed: {video.get('created_at', 'N/A')}",
            f"Analysis Type: {video.get('analysis_type', 'summarize')}",
            "",
            "TRANSCRIPTION",
            "-" * 50,
            "",
            video.get('transcription', ''),
            "",
            "ANALYSIS",
            "-" * 50,
            "",
            video.get('analysis', ''),
        ]
        return "\n".join(lines)
    
    def export_pdf(self, video):
        """Export video to PDF format"""
        from io import BytesIO
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#667eea',
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='#764ba2',
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        story.append(Paragraph("Video Transcription & Analysis", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        metadata = [
            f"<b>Video URL:</b> {video.get('video_url', 'N/A')}",
            f"<b>Platform:</b> {video.get('platform', 'unknown')}",
            f"<b>Duration:</b> {video.get('duration_minutes', 0):.1f} minutes",
            f"<b>Processed:</b> {video.get('created_at', 'N/A')}",
            f"<b>Analysis Type:</b> {video.get('analysis_type', 'summarize')}",
        ]
        
        for meta in metadata:
            story.append(Paragraph(meta, styles['Normal']))
        
        story.append(Spacer(1, 0.3*inch))
        
        # Transcription
        story.append(Paragraph("TRANSCRIPTION", heading_style))
        transcription = video.get('transcription', '').replace('\n', '<br/>')
        story.append(Paragraph(transcription, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Analysis
        story.append(Paragraph("ANALYSIS", heading_style))
        analysis = video.get('analysis', '').replace('\n', '<br/>')
        story.append(Paragraph(analysis, styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_docx(self, video):
        """Export video to DOCX format"""
        from io import BytesIO
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Video Transcription & Analysis', 0)
        title.alignment = 1  # Center
        
        # Metadata
        doc.add_paragraph(f"Video URL: {video.get('video_url', 'N/A')}")
        doc.add_paragraph(f"Platform: {video.get('platform', 'unknown')}")
        doc.add_paragraph(f"Duration: {video.get('duration_minutes', 0):.1f} minutes")
        doc.add_paragraph(f"Processed: {video.get('created_at', 'N/A')}")
        doc.add_paragraph(f"Analysis Type: {video.get('analysis_type', 'summarize')}")
        
        doc.add_paragraph()  # Blank line
        
        # Transcription
        doc.add_heading('TRANSCRIPTION', level=1)
        doc.add_paragraph(video.get('transcription', ''))
        
        doc.add_paragraph()  # Blank line
        
        # Analysis
        doc.add_heading('ANALYSIS', level=1)
        doc.add_paragraph(video.get('analysis', ''))
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

